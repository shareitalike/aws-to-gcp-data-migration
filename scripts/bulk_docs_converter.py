import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Add a professional header placeholder for the user
    # header = doc.sections[0].header
    # header.paragraphs[0].text = "[YOUR NAME] | [YOUR PHONE] | [YOUR EMAIL]"
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    
    for line in lines:
        line = line.strip()
        
        # Handle Code Blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            if line:
                p = doc.add_paragraph(line)
                p.style = 'No Spacing'
                if p.runs:
                    run = p.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            continue

        # Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Handle Bulllets
        elif line.startswith('* ') or line.startswith('- '):
            content = line[2:]
            # basic bold parsing
            content = content.replace('**', '') 
            doc.add_paragraph(content, style='List Bullet')
            
        # Handle Regular Text
        elif line:
            # Simple bold detection/stripping for cleaner doc
            content = line.replace('**', '')
            doc.add_paragraph(content)

    doc.save(docx_path)
    print(f"  Converted: {os.path.basename(md_path)} -> {os.path.basename(docx_path)}")

def main():
    root_dir = Path("f:/pyspark_study/project_bigquery_auto")
    docs_dir = root_dir / "docs"
    export_dir = docs_dir / "word_exports"
    
    if not export_dir.exists():
        export_dir.mkdir(parents=True)
        
    print(f"Starting bulk conversion in: {docs_dir}")
    
    # Files in docs root
    for md_file in docs_dir.glob("*.md"):
        docx_name = md_file.stem + ".docx"
        convert_md_to_docx(md_file, export_dir / docx_name)
        
    # Files in interview_prep subfolder
    prep_dir = docs_dir / "interview_prep"
    if prep_dir.exists():
        for md_file in prep_dir.glob("*.md"):
            docx_name = md_file.stem + ".docx"
            convert_md_to_docx(md_file, export_dir / docx_name)
            
    # Also handle README.md at root
    readme_md = root_dir / "README.md"
    if readme_md.exists():
        convert_md_to_docx(readme_md, export_dir / "PROJECT_README.docx")

    print("\nBulk conversion complete! Check the 'docs/word_exports' folder.")

if __name__ == "__main__":
    main()
